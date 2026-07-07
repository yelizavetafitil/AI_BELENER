import os
import io
import re
import base64
import json
import logging
import time
import hmac
import hashlib
import tempfile
import uuid
import threading
import urllib.request
from flask import Flask, request, Response, send_from_directory, stream_with_context, session, redirect, jsonify
import psycopg2
import psycopg2.extras
import ollama

from belener.config import model_drawing, model_scan, report_llm_enabled, ensure_upload_temp_dir
from belener.extract import extract_pdf_path
from belener.extract_report import extraction_to_markdown
from belener.scanned import is_scanned_pdf

app = Flask(__name__, static_folder="static")
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s %(message)s")
logging.getLogger("belener").setLevel(logging.INFO)
ROOT_DIR = app.root_path
app.secret_key = os.environ.get("AI_FLASK_SECRET", "change-this-ai-flask-secret")
app.config.update(
    SESSION_COOKIE_NAME='belnipiai_session',
    SESSION_COOKIE_PATH='/',
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax'
)

MODEL_DEFAULT = os.environ.get("MODEL_DEFAULT", "gemma3:4b")
AI_SSO_SHARED_SECRET = os.environ.get("AI_SSO_SHARED_SECRET", "change-this-ai-sso-secret")
AI_SSO_MAX_AGE_SEC = int(os.environ.get("AI_SSO_MAX_AGE_SEC", "300"))

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
DATABASE_URL = os.environ.get("DATABASE_URL", "postgresql://postgres:sdd05072008sdd@localhost:5432/belnipiai")

_client = ollama.Client(host=OLLAMA_HOST)


def _normalize_username(raw_username: str) -> str:
    login = (raw_username or "").strip().lower()
    if "\\" in login:
        login = login.split("\\")[-1].strip()
    if "@" in login:
        login = login.split("@")[0].strip()
    return login


def _verify_sso_payload(username: str, display_name: str, ts_raw: str, sig: str) -> bool:
    try:
        ts_val = int(ts_raw)
    except Exception:
        return False
    if abs(int(time.time()) - ts_val) > AI_SSO_MAX_AGE_SEC:
        return False
    payload = f"{username}|{display_name}|{ts_raw}"
    expected = hmac.new(
        AI_SSO_SHARED_SECRET.encode("utf-8"),
        payload.encode("utf-8"),
        hashlib.sha256
    ).hexdigest()
    return hmac.compare_digest(expected, sig or "")


# ── модели ────────────────────────────────────────────────────────────────────

# Отображаемые названия для фронтенда
MODEL_LABELS = {
    "gemma3:4b":  "Базовая",   # ноутбук
    "gemma4:e4b": "Базовая",   # сервер
    "gemma4:26b": "Pro",        # сервер
}

MODEL_PROFILES = {
    "gemma3:4b":  {"num_ctx": 14000, "chunk_tokens": 11000, "no_think": False},
    "gemma4:e4b": {"num_ctx": 14000, "chunk_tokens": 11000, "no_think": False},
    "gemma4:26b": {"num_ctx": 14000, "chunk_tokens": 11000, "no_think": False},
}

DEFAULT_PROFILE = {"num_ctx": 14000, "chunk_tokens": 11000, "no_think": False}

MODEL_TASK_PRIORITY = {
    "image":    (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Изображение — нужна vision-модель"),
    "document": (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Документ — нужен большой контекст"),
    "table":    (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Таблица — нужна аналитическая модель"),
    "complex":  (["gemma4:26b", "gemma4:e4b", "gemma3:4b"], "Сложный вопрос — нужна умная модель"),
    "simple":   (["gemma3:4b", "gemma4:e4b", "gemma4:26b"], "Простой вопрос"),
}

COMPLEX_KEYWORDS = {
    "проанализируй", "сравни", "объясни", "почему", "как работает",
    "разбери", "оцени", "составь", "напиши", "рассчитай", "найди противоречия",
}


def _resolve_config_model(name: str, available: list) -> str | None:
    """Модель из .env, если она есть в Ollama."""
    if not name:
        return None
    if name in available:
        return name
    base = name.split(":")[0]
    for m in available:
        if m == name or m.startswith(base + ":") or m.startswith(base):
            return m
    return name


_VISION_MODEL_HINTS = ("vl", "vision", "llava", "moondream", "minicpm-v", "bakllava")


def is_chat_model(model_id: str) -> bool:
    """Модели для выпадающего списка чата (без vision — зоны в PDF_VISION_MODEL)."""
    if os.environ.get("MODEL_CHAT_INCLUDE_VISION", "").strip().lower() in ("1", "true", "yes", "on"):
        return True
    low = model_id.lower()
    return not any(h in low for h in _VISION_MODEL_HINTS)


def model_for_task(task: str, user_model: str, *, user_override: bool = False) -> str:
    """task: drawing | scan — MODEL_DRAWING / MODEL_SCAN, если пользователь не выбрал модель вручную."""
    if not user_override:
        env_name = model_drawing() if task == "drawing" else model_scan() if task == "scan" else ""
        if env_name:
            available = get_available_models()
            resolved = _resolve_config_model(env_name, available)
            if resolved:
                return resolved
    return user_model or MODEL_DEFAULT


def suggest_model(file_ext: str, question: str, available: list) -> tuple[str, str]:
    """Возвращает (model_name, reason) из доступных моделей."""

    def pick(task_key: str) -> str | None:
        preferred, _ = MODEL_TASK_PRIORITY[task_key]
        for p in preferred:
            base = p.split(":")[0]
            for m in available:
                if m == p or m.startswith(base):
                    return m
        return None

    images = {".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"}
    docs   = {".pdf", ".docx", ".doc"}
    tables = {".xlsx", ".xls", ".xlsm", ".ods", ".csv", ".tsv"}

    if file_ext in images:
        task = "image"
    elif file_ext in docs:
        task = "document"
    elif file_ext in tables:
        task = "table"
    elif len(question) > 180 or any(kw in question.lower() for kw in COMPLEX_KEYWORDS):
        task = "complex"
    else:
        task = "simple"

    model = pick(task) or (available[0] if available else MODEL_DEFAULT)
    _, reason = MODEL_TASK_PRIORITY[task]
    return model, reason


SYSTEM_PROMPT_DOC = (
    "Ты аналитик документов. Отвечай только на русском языке. "
    "Используй только предоставленный текст документа. "
    "Не переводи на английский. Не пересказывай и не сокращай: "
    "если просят извлечь или прочитать текст — выводи его полностью, как в документе. "
    "Не выдумывай факты, организации и номера. Формат ответа: Markdown."
)

SYSTEM_PROMPT_CHAT = (
    "You are a helpful assistant. "
    "If you are not certain about something, say so — never make up facts, numbers or dates. "
    "Format your responses in Markdown."
)

EXTRACT_DEFAULT_QUESTION = "Извлечь весь текст с листа"
GOST_DEFAULT_QUESTION = "Проверка ГОСТ на листе"
DRAWING_PROCESS_STATUS = "Обрабатываю чертеж…"
EXTRACT_FOLLOWUP_PROMPT = (
    "Ниже — **только** извлечённый с листа текст (OCR/текстовый слой). "
    "Отвечай **строго** по нему на русском. Не придумывай факты.\n\n"
)


def _is_default_drawing_question(q: str) -> bool:
    s = (q or "").strip().casefold()
    return s in {
        EXTRACT_DEFAULT_QUESTION.casefold(),
        "разбор документа",
        "анализ текста",
        "разбор листа",
        "извлечь весь текст",
        "извлечь текст",
        "извлеки текст",
        "извлеки текст с листа",
        "извлечь текст с листа",
        "текст с листа",
        "прочитай скан",
        "прочитай лист",
        "разбор документа и анализ текста",
    }


def _parse_check_date(raw: str | None):
    """Дата проверки актуальности нормативов (по умолчанию — сегодня на сервере)."""
    from datetime import date, datetime

    s = (raw or "").strip()
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    return None


def _is_gost_check_request(question: str, mode: str = "") -> bool:
    if _is_text_extract_question(question):
        return False
    if (mode or "").strip().casefold() in ("gost", "1", "true", "yes", "on"):
        return True
    return _is_normative_question(question)


def _is_normative_question(q: str) -> bool:
    """Запрос на перечень ГОСТ/нормативов — быстрый OCR, без полного разбора чертежа."""
    s = (q or "").strip().casefold().replace("ё", "е")
    if not s:
        return False
    keys = (
        GOST_DEFAULT_QUESTION.casefold(),
        "проверка гост",
        "проверить гост",
        "все гост",
        "все gost",
        "все ost",
        "все ост",
        "все норматив",
        "найди гост",
        "найти гост",
        "найди gost",
        "список гост",
        "список gost",
        "какие гост",
        "перечень гост",
        "гост на листе",
        "нормативн",
    )
    return any(k in s for k in keys)


def _is_analysis_question(q: str) -> bool:
    s = (q or "").strip().casefold().replace("ё", "е")
    return any(k in s for k in ("разбор", "анализ"))


def _is_text_extract_question(q: str) -> bool:
    s = (q or "").strip().casefold().replace("ё", "е")
    return any(
        k in s
        for k in (
            "извлеки текст",
            "извлечь текст",
            "весь текст",
            "прочитай лист",
            "прочитай скан",
            "текст с листа",
        )
    )


def _is_full_extract_question(q: str) -> bool:
    if _is_normative_question(q):
        return True
    return _is_default_drawing_question(q) or not (q or "").strip()


def _chunk_sse_text(text: str, size: int = 900):
    for i in range(0, len(text), size):
        yield f"data: {json.dumps({'text': text[i:i + size]}, ensure_ascii=False)}\n\n"


def _ollama_user_message(exc: Exception) -> str:
    msg = str(exc).strip()
    low = msg.lower()
    if "timed out" in low or "timeout" in low:
        return (
            "Таймаут при обращении к normy.stn.by. Список ГОСТ на листе сохранён; "
            "проверку актуальности повторите позже или увеличьте PDF_STN_TIMEOUT в .env."
        )
    if "model runner" in low or "unexpectedly stopped" in low:
        return (
            "Vision-модель Ollama остановилась (часто нехватка RAM на CPU). "
            "Перезапустите Ollama или используйте текстовую модель; для чертежей — зонный OCR."
        )
    if "connection" in low or "refused" in low:
        return f"Нет связи с Ollama ({OLLAMA_HOST}). Проверьте контейнер ollama."
    return msg or exc.__class__.__name__


def _sse_error(msg: str):
    yield f"data: {json.dumps({'error': msg}, ensure_ascii=False)}\n\n"
    yield "data: [DONE]\n\n"


def _sse_text(text: str):
    """Статус/фрагмент в чат (без эмодзи, нормальные переносы строк)."""
    yield f"data: {json.dumps({'text': text}, ensure_ascii=False)}\n\n"


def _sse_status(text: str):
    yield f"data: {json.dumps({'status': text}, ensure_ascii=False)}\n\n"


def stream_extract_pdf_normative(path: str, filename: str, question: str, *, check_date=None):
    """PDF → OCR страниц (плитки) → таблица нормативов + проверка ГОСТ на STN."""
    import time
    from datetime import date

    from belener.normative_extract import extract_normatives_pdf_path, normative_result_to_markdown

    validity_date = check_date or date.today()

    yield from _sse_status(DRAWING_PROCESS_STATUS)

    box: dict = {"result": None, "err": None}
    done = threading.Event()

    def _run():
        try:
            box["result"] = extract_normatives_pdf_path(path, filename)
        except Exception as e:
            app.logger.exception("normative PDF extract failed file=%s", filename)
            box["err"] = e
        finally:
            done.set()

    threading.Thread(target=_run, daemon=True).start()
    while not done.wait(timeout=12):
        yield from _sse_status(DRAWING_PROCESS_STATUS)

    if box["err"] is not None:
        yield from _sse_error(f"Ошибка: {_ollama_user_message(box['err'])}")
        return

    result = box["result"] or {}
    if not result.get("ok"):
        yield from _sse_error(str(result.get("error") or "Не удалось прочитать PDF"))
        return

    from belener.config import gost_check_total_budget_sec, stn_batch_budget_sec, stn_lookup_enabled
    from belener.stn_lookup import refine_and_check_normative_refs

    stn_checks = []
    if stn_lookup_enabled():
        yield from _sse_status(DRAWING_PROCESS_STATUS)
        try:
            page_count = int(result.get("page_count") or 1)
            # Резерв STN после OCR: не привязывать к старту pipeline (OCR ~100+ с).
            stn_deadline = time.monotonic() + stn_batch_budget_sec()
            refined, stn_checks = refine_and_check_normative_refs(
                result.get("normative_refs") or [],
                today=validity_date,
                deadline=stn_deadline,
            )
            result["normative_refs"] = refined
            result["stn_checks"] = [c.to_dict() for c in stn_checks]
        except Exception as e:
            app.logger.exception("STN batch failed file=%s", filename)
            result["stn_error"] = _ollama_user_message(e)

    include_ctx = "контекст" in (question or "").casefold()
    report = normative_result_to_markdown(
        result,
        include_context=include_ctx,
        stn_checks=stn_checks,
        check_date=validity_date,
    )
    yield from _chunk_sse_text(report)
    yield "data: [DONE]\n\n"


def stream_extract_image_normative(path: str, filename: str, question: str, *, check_date=None):
    """Изображение → OCR → таблица нормативов."""
    import time
    from datetime import date

    from belener.normative_extract import extract_normatives_from_image_path, normative_result_to_markdown

    validity_date = check_date or date.today()

    yield from _sse_status(DRAWING_PROCESS_STATUS)
    try:
        result = extract_normatives_from_image_path(path, filename)
    except Exception as e:
        app.logger.exception("normative image OCR failed file=%s", filename)
        yield from _sse_error(f"Ошибка OCR: {_ollama_user_message(e)}")
        return

    include_ctx = "контекст" in (question or "").casefold()
    from belener.config import stn_batch_budget_sec, stn_lookup_enabled
    from belener.stn_lookup import refine_and_check_normative_refs

    stn_checks = []
    if stn_lookup_enabled():
        yield from _sse_status(DRAWING_PROCESS_STATUS)
        try:
            stn_deadline = time.monotonic() + stn_batch_budget_sec()
            refined, stn_checks = refine_and_check_normative_refs(
                result.get("normative_refs") or [],
                today=validity_date,
                deadline=stn_deadline,
            )
            result["normative_refs"] = refined
            result["stn_checks"] = [c.to_dict() for c in stn_checks]
        except Exception as e:
            app.logger.exception("STN batch failed file=%s", filename)
            result["stn_error"] = _ollama_user_message(e)
    report = normative_result_to_markdown(
        result,
        include_context=include_ctx,
        stn_checks=stn_checks,
        check_date=validity_date,
    )
    yield from _chunk_sse_text(report)
    yield "data: [DONE]\n\n"


def _skip_extract_followup(question: str) -> bool:
    """Не дублировать отчёт в «Ответ по вопросу»."""
    if _is_full_extract_question(question) or _is_text_extract_question(question) or _is_analysis_question(question):
        return True
    q = (question or "").strip()
    if len(q) < 30:
        return True
    low = q.casefold()
    if low in ("copy", "копия", "скопируй", "вывод", "отчёт", "отчет"):
        return True
    return False


def stream_extract_pdf(path: str, filename: str, question: str, model: str, history=None):
    """Извлечение листа: tile OCR (тот же путь, что для ГОСТ)."""
    yield from _sse_status(DRAWING_PROCESS_STATUS)

    box: dict = {"facts": None, "body": None, "err": None}
    ocr_done = threading.Event()
    fmt_done = threading.Event()

    def _run_ocr():
        try:
            box["facts"] = extract_pdf_path(path, filename)
        except Exception as e:
            app.logger.exception("PDF extract failed file=%s", filename)
            box["err"] = e
        finally:
            ocr_done.set()

    def _run_format():
        ocr_done.wait()
        if box["err"] is not None or not (box["facts"] or {}).get("ok"):
            fmt_done.set()
            return
        try:
            box["body"] = extraction_to_markdown(box["facts"], question=question, polish_llm=True)
        except Exception as e:
            app.logger.exception("report format failed file=%s", filename)
            box["err"] = e
        finally:
            fmt_done.set()

    threading.Thread(target=_run_ocr, daemon=True).start()
    threading.Thread(target=_run_format, daemon=True).start()
    while not fmt_done.wait(timeout=12):
        yield from _sse_status(DRAWING_PROCESS_STATUS)

    if box["err"] is not None:
        yield from _sse_error(f"Ошибка извлечения: {_ollama_user_message(box['err'])}")
        return
    facts = box["facts"]
    if not facts or not facts.get("ok"):
        yield from _sse_error(str((facts or {}).get("error") or "Не удалось извлечь текст"))
        return

    body = box["body"] or extraction_to_markdown(facts, question=question, polish_llm=True)
    report = f"**Файл:** {filename}\n\n{body}"
    yield from _chunk_sse_text(report)

    if _skip_extract_followup(question):
        yield "data: [DONE]\n\n"
        return

    yield from _sse_text("\n\n---\n\n## Ответ по вашему вопросу\n\n")
    follow_q = (
        EXTRACT_FOLLOWUP_PROMPT
        + f"<extracted_text>\n{report}\n</extracted_text>\n\nВопрос: {question}"
    )
    profile = get_profile(model)
    messages = [{"role": "system", "content": SYSTEM_PROMPT_DOC}]
    if history:
        messages += history
    messages.append({"role": "user", "content": apply_no_think(follow_q, model)})
    try:
        for chunk in _client.chat(
            model=model,
            messages=messages,
            options={"num_ctx": profile["num_ctx"], "temperature": 0.1},
            stream=True,
        ):
            delta = chunk["message"]["content"]
            yield f"data: {json.dumps({'text': delta}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"
    except Exception as e:
        yield from _sse_error(f"Ошибка ответа: {_ollama_user_message(e)}")


def get_profile(model: str) -> dict:
    if model in MODEL_PROFILES:
        return MODEL_PROFILES[model]
    for key, profile in MODEL_PROFILES.items():
        if model.startswith(key.split(":")[0]):
            return profile
    return DEFAULT_PROFILE


def apply_no_think(text: str, model: str) -> str:
    if get_profile(model)["no_think"]:
        return "/no_think " + text
    return text


def estimate_tokens(text: str) -> int:
    """Быстрая локальная оценка без сетевого вызова (для решений о чанкинге)."""
    cyrillic = sum(1 for c in text if 'Ѐ' <= c <= 'ӿ')
    return max(1, cyrillic // 2 + (len(text) - cyrillic) // 3)


def count_tokens(text: str, model: str = MODEL_DEFAULT) -> int:
    """Точный подсчёт через Ollama /api/tokenize, fallback на estimate_tokens."""
    try:
        body = json.dumps({"model": model, "prompt": text}).encode()
        req = urllib.request.Request(
            f"{OLLAMA_HOST}/api/tokenize",
            data=body,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=3) as resp:
            return max(1, len(json.loads(resp.read()).get("tokens", [])))
    except Exception:
        return estimate_tokens(text)


# ── база данных ───────────────────────────────────────────────────────────────

def get_db():
    conn = psycopg2.connect(DATABASE_URL, cursor_factory=psycopg2.extras.RealDictCursor)
    return conn


def init_db():
    schema = os.path.join(ROOT_DIR, "schema.sql")
    if not os.path.exists(schema):
        return
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(open(schema).read())
        conn.commit()


# ── управление контекстом ─────────────────────────────────────────────────────

def load_history(conv_id: str) -> list:
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT role, content FROM messages "
                "WHERE conversation_id = %s AND role != 'summary' "
                "ORDER BY created_at",
                (conv_id,)
            )
            rows = cur.fetchall()
    return [{"role": r["role"], "content": r["content"]} for r in rows]


def load_history_with_summary(conv_id: str, num_ctx: int) -> list:
    """Загружает историю. Для сообщений с файлами восстанавливает документ из БД."""
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT content, created_at FROM messages "
                "WHERE conversation_id = %s AND role = 'summary' "
                "ORDER BY created_at DESC LIMIT 1",
                (conv_id,)
            )
            summary_row = cur.fetchone()

            since = summary_row["created_at"] if summary_row else None
            base_q = """
                SELECT m.id, m.role, m.content, m.token_count,
                       a.extracted_text, a.original_name
                FROM messages m
                LEFT JOIN attachments a ON a.message_id = m.id
                WHERE m.conversation_id = %s AND m.role != 'summary'
            """
            if since:
                cur.execute(base_q + " AND m.created_at > %s ORDER BY m.created_at", (conv_id, since))
            else:
                cur.execute(base_q + " ORDER BY m.created_at", (conv_id,))
            rows = cur.fetchall()

    rows = list(rows)

    # Только последний документ получает полный текст в контекст,
    # старые документы — лишь ссылка на имя файла
    max_doc_chars = int(num_ctx * 0.55) * 4
    last_doc_idx = max(
        (i for i, r in enumerate(rows) if r["role"] == "user" and r["extracted_text"]),
        default=None
    )

    def build_content(row, idx):
        if row["role"] == "user" and row["extracted_text"]:
            if idx == last_doc_idx:
                doc = row["extracted_text"]
                if len(doc) > max_doc_chars:
                    doc = doc[:max_doc_chars] + "\n\n[... документ обрезан из-за лимита контекста ...]"
                return f"<document name=\"{row['original_name']}\">\n{doc}\n</document>\n\nВопрос: {row['content']}"
            return f"[Ранее загружен документ: {row['original_name']}]\n\nВопрос: {row['content']}"
        return row["content"]

    messages = []
    if summary_row:
        messages.append({"role": "user", "content": f"[Резюме предыдущего разговора]\n{summary_row['content']}"})
        messages.append({"role": "assistant", "content": "Понял, продолжаю разговор с учётом предыдущего контекста."})

    # Считаем токены по фактическому контенту (с документом), а не по r["content"].
    # Это правильно при смене модели: сохранённый token_count от старой модели не используется.
    total_tokens = sum(estimate_tokens(build_content(r, i)) for i, r in enumerate(rows))
    threshold = int(num_ctx * 0.85)

    if total_tokens <= threshold or len(rows) <= 4:
        messages += [{"role": r["role"], "content": build_content(r, i)} for i, r in enumerate(rows)]
        return messages

    # Нужна обрезка. Бюджет = threshold минус уже добавленные summary-сообщения.
    budget = threshold - sum(estimate_tokens(m["content"]) for m in messages)

    # Якоря: последний документ + ответ на него — всегда включаем, независимо от бюджета.
    anchors: set[int] = set()
    if last_doc_idx is not None:
        anchors.add(last_doc_idx)
        nxt = last_doc_idx + 1
        if nxt < len(rows) and rows[nxt]["role"] == "assistant":
            anchors.add(nxt)

    for idx in sorted(anchors):
        budget -= estimate_tokens(build_content(rows[idx], idx))

    # Добираем свежие сообщения (от конца) в рамках оставшегося бюджета.
    tail: set[int] = set()
    for i in range(len(rows) - 1, -1, -1):
        if i in anchors:
            continue
        t = estimate_tokens(rows[i]["content"])
        if budget - t < 100:
            break
        tail.add(i)
        budget -= t

    keep = sorted(anchors | tail)
    messages += [{"role": rows[i]["role"], "content": build_content(rows[i], i)} for i in keep]

    return messages


def maybe_summarize(conv_id: str, model: str, num_ctx: int):
    """Если история занимает > 65% контекста — делаем резюме (вызывается из фонового треда)."""
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT SUM(token_count) FROM messages "
                "WHERE conversation_id = %s AND role != 'summary'",
                (conv_id,)
            )
            total = cur.fetchone()["sum"] or 0

    if total < int(num_ctx * 0.65):
        return

    history = load_history(conv_id)
    if len(history) < 6:
        return

    old_messages = history[:-4]
    text_to_summarize = "\n".join(f"{m['role'].upper()}: {m['content']}" for m in old_messages)

    resp = _client.chat(
        model=model,
        messages=[{
            "role": "user",
            "content": f"Сделай краткое резюме этого диалога (важные факты, решения, контекст):\n\n{text_to_summarize}"
        }],
        options={"num_ctx": num_ctx, "temperature": 0.1},
        stream=False
    )
    summary_text = resp["message"]["content"]

    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (conversation_id, role, content, token_count) VALUES (%s, 'summary', %s, %s)",
                (conv_id, summary_text, count_tokens(summary_text, model))
            )
        conn.commit()


def save_message(conv_id: str, role: str, content: str, model: str = MODEL_DEFAULT,
                 attachment_name: str = None, attachment_text: str = None):
    tokens = count_tokens(content, model)
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO messages (id, conversation_id, role, content, token_count) VALUES (%s, %s, %s, %s, %s) RETURNING id",
                (str(uuid.uuid4()), conv_id, role, content, tokens)
            )
            msg_id = cur.fetchone()["id"]
            if attachment_name:
                cur.execute(
                    "INSERT INTO attachments (message_id, original_name, extracted_text) VALUES (%s, %s, %s)",
                    (msg_id, attachment_name, attachment_text or "")
                )
            cur.execute(
                "UPDATE conversations SET updated_at = NOW() WHERE id = %s",
                (conv_id,)
            )
        conn.commit()


def derive_chat_title(question: str, file_name: str | None = None) -> str:
    """Название чата по первому сообщению (без LLM — без «Гости под звездами» и т.п.)."""
    q = re.sub(r"\s+", " ", (question or "").strip())
    q_one = q.split("\n")[0].strip()

    if _is_normative_question(q):
        task = "ГОСТ"
    elif _is_text_extract_question(q) or _is_default_drawing_question(q):
        task = "Разбор листа"
    elif _is_analysis_question(q):
        task = "Анализ"
    else:
        task = ""

    stem = ""
    if file_name:
        stem = os.path.splitext(os.path.basename(file_name))[0].strip()

    if stem and task:
        title = f"{task}: {stem}"
    elif stem:
        title = stem
    elif task:
        title = task
    elif q_one:
        title = q_one
    else:
        title = "Новый чат"

    title = title.strip('"\'«»').rstrip(".,;:!?").strip()
    if len(title) > 60:
        if stem and task:
            budget = max(8, 60 - len(task) - 2)
            title = f"{task}: {stem[:budget].rstrip()}"
        else:
            title = title[:57].rstrip() + "…"
    return title if len(title) >= 2 else "Новый чат"


# ── утилиты файлов ────────────────────────────────────────────────────────────

def to_base64(pil_image):
    buf = io.BytesIO()
    pil_image.save(buf, format="PNG", optimize=False, compress_level=0)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def get_available_models():
    try:
        result = _client.list()
        return [m.model for m in result.models]
    except Exception:
        return [MODEL_DEFAULT]


def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def detect_pdf_type(path):
    import pymupdf
    doc = pymupdf.open(path)
    total_chars = sum(len(p.get_text().strip()) for p in doc)
    avg = total_chars / max(len(doc), 1)
    return "scanned" if avg < 50 else "text"


def read_pdf_text(path):
    import pymupdf
    doc = pymupdf.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"[Страница {i+1}]\n{text}")
    return pages


def read_docx(path):
    import docx
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_excel(path):
    import pandas as pd
    ext = os.path.splitext(path)[1].lower()
    xf = pd.ExcelFile(path, engine="xlrd" if ext == ".xls" else "openpyxl")
    parts = []
    for sheet in xf.sheet_names:
        df = xf.parse(sheet)
        parts.append(f"[Лист: {sheet}]\n{df.to_string()}")
    return "\n\n".join(parts)


def read_csv(path, ext=".csv"):
    import pandas as pd
    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(path, sep=sep, encoding="utf-8", errors="ignore")
    except Exception:
        df = pd.read_csv(path, sep=sep, encoding="cp1251", errors="ignore")
    return df.to_string()


# ── стриминг ──────────────────────────────────────────────────────────────────

def stream_text_response(content, question, model, history=None):
    profile = get_profile(model)
    num_ctx = profile["num_ctx"]

    max_chars = (num_ctx - 1000) * 4
    if len(content) > max_chars:
        content = content[:max_chars]

    messages = [{"role": "system", "content": SYSTEM_PROMPT_DOC}]
    if history:
        messages += history
    messages.append({
        "role": "user",
        "content": apply_no_think(f"<document>\n{content}\n</document>\n\nВопрос: {question}", model)
    })

    try:
        response = _client.chat(
            model=model,
            messages=messages,
            options={"num_ctx": num_ctx, "temperature": 0.1},
            stream=True,
        )
        for chunk in response:
            yield f"data: {json.dumps({'text': chunk['message']['content']})}\n\n"
        yield "data: [DONE]\n\n"
    except Exception as e:
        yield from _sse_error(_ollama_user_message(e))


def stream_image_response(images_b64, question, model, history=None):
    profile = get_profile(model)
    messages = []
    if history:
        messages += history
    messages.append({
        "role": "user",
        "content": apply_no_think(question, model),
        "images": images_b64,
    })
    try:
        response = _client.chat(
            model=model,
            messages=messages,
            options={"num_ctx": profile["num_ctx"], "temperature": 0.1},
            stream=True,
        )
        for chunk in response:
            yield f"data: {json.dumps({'text': chunk['message']['content']})}\n\n"
        yield "data: [DONE]\n\n"
    except Exception as e:
        yield from _sse_error(_ollama_user_message(e))


def stream_map_reduce(pages_text, question, model, history=None):
    profile = get_profile(model)
    num_ctx = profile["num_ctx"]
    chunk_chars = profile["chunk_tokens"] * 4

    chunks, current, cur_len = [], [], 0
    for page in pages_text:
        if cur_len + len(page) > chunk_chars and current:
            chunks.append("\n\n".join(current))
            current, cur_len = [page], len(page)
        else:
            current.append(page)
            cur_len += len(page)
    if current:
        chunks.append("\n\n".join(current))

    total = len(chunks)
    yield from _sse_text(f"Документ большой — обрабатываю {total} частей...\n\n")

    summaries, prev_context = [], ""
    for i, chunk in enumerate(chunks):
        yield from _sse_text(f"Часть {i + 1} из {total}...\n")
        context_hint = f"\n\nКонтекст из предыдущих частей:\n{prev_context}\n" if prev_context else ""
        prompt = apply_no_think(
            f"Это часть {i+1} из {total} документа.\nВопрос: {question}{context_hint}\n"
            f"Извлеки всё относящееся к вопросу, сохрани детали, цифры, имена:\n\n{chunk}", model
        )
        s = ""
        for chunk_r in _client.chat(model=model,
                messages=[{"role": "user", "content": prompt}],
                                     options={"num_ctx": num_ctx, "temperature": 0.1}, stream=True):
            s += chunk_r["message"]["content"]
        summaries.append(f"[Часть {i+1}/{total}]\n{s}")
        prev_context = s[-500:] if len(s) > 500 else s

    yield from _sse_text("\n\n---\n\n## Финальный ответ\n\n")
    yield from stream_text_response("\n\n".join(summaries), question, model, history)


# ── роуты: конверсации ────────────────────────────────────────────────────────

@app.route("/api/conversations", methods=["GET"])
def api_list_conversations():
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id, title, model, updated_at FROM conversations ORDER BY updated_at DESC")
            rows = cur.fetchall()
    return {"conversations": [dict(r) for r in rows]}


@app.route("/api/conversations", methods=["POST"])
def api_create_conversation():
    data = request.get_json(force=True)
    model = data.get("model", MODEL_DEFAULT)
    cid = str(uuid.uuid4())
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO conversations (id, title, model) VALUES (%s, %s, %s)",
                (cid, "Новый чат", model)
            )
        conn.commit()
    return {"id": cid, "title": "Новый чат", "model": model}


@app.route("/api/conversations/<conv_id>", methods=["DELETE"])
def api_delete_conversation(conv_id):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM conversations WHERE id = %s", (conv_id,))
        conn.commit()
    return {"ok": True}


@app.route("/api/conversations/<conv_id>/messages", methods=["GET"])
def api_get_messages(conv_id):
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT m.id, m.role, m.content, m.created_at, "
                "a.original_name as file_name "
                "FROM messages m "
                "LEFT JOIN attachments a ON a.message_id = m.id "
                "WHERE m.conversation_id = %s AND m.role != 'summary' "
                "ORDER BY m.created_at",
                (conv_id,)
            )
            rows = cur.fetchall()
    return {"messages": [dict(r) for r in rows]}


@app.route("/api/conversations/<conv_id>/title", methods=["PATCH"])
def api_update_title(conv_id):
    data = request.get_json(force=True)
    title = data.get("title", "")[:80]
    with get_db() as conn:
        with conn.cursor() as cur:
            cur.execute("UPDATE conversations SET title = %s WHERE id = %s", (title, conv_id))
        conn.commit()
    return {"ok": True}


# ── роут: чат ─────────────────────────────────────────────────────────────────

@app.route("/api/conversations/<conv_id>/chat", methods=["POST"])
def api_chat(conv_id):
    question = request.form.get("question", "").strip()
    model = request.form.get("model", MODEL_DEFAULT)
    user_override = request.form.get("model_override", "").strip().lower() in ("1", "true", "yes", "on")
    gost_mode = request.form.get("mode", "").strip()
    check_date = _parse_check_date(request.form.get("check_date"))
    file = request.files.get("file")

    tmp_path = file_name = extracted_text = None
    ext = ""
    if file and file.filename:
        ext = os.path.splitext(file.filename)[1].lower()
        file_name = file.filename
        upload_dir = ensure_upload_temp_dir()
        tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False, dir=upload_dir)
        tmp_path = tmp.name
        tmp.close()
        try:
            file.save(tmp_path)
        except OSError as e:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            if getattr(e, "errno", None) == 28:
                return {"error": "Недостаточно места на диске для загрузки файла. Очистите кэш или освободите место на диске D:."}, 507
            raise

    if not question:
        if tmp_path and ext in (".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"):
            question = GOST_DEFAULT_QUESTION
        else:
            return {"error": "Вопрос не может быть пустым"}, 400

    profile = get_profile(model)

    # Файл уже сохранён выше

    # Загружаем историю для контекста
    history = load_history_with_summary(conv_id, profile["num_ctx"])

    # Сохраняем вопрос пользователя
    save_message(
        conv_id,
        "user",
        question,
        model,
        attachment_name=file_name if file_name else None,
    )

    def generate():
        nonlocal extracted_text
        full_response = []

        try:
            if not tmp_path:
                # Обычный чат с историей
                messages = [{"role": "system", "content": SYSTEM_PROMPT_CHAT}]
                messages += history
                messages.append({"role": "user", "content": apply_no_think(question, model)})

                for chunk in _client.chat(
                    model=model,
                    messages=messages,
                    options={"num_ctx": profile["num_ctx"], "temperature": 0.7},
                    stream=True,
                ):
                    delta = chunk["message"]["content"]
                    full_response.append(delta)
                    yield f"data: {json.dumps({'text': delta})}\n\n"
                yield "data: [DONE]\n\n"

            else:
                if ext in (".txt", ".md", ".log"):
                    text = read_txt(tmp_path)
                    extracted_text = text
                    gen = (stream_text_response(text, question, model, history)
                           if estimate_tokens(text) <= profile["chunk_tokens"]
                           else stream_map_reduce(
                               [text[i:i+profile["chunk_tokens"]*4]
                                for i in range(0, len(text), profile["chunk_tokens"]*4)],
                               question, model, history))
                    for chunk in gen:
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext == ".pdf":
                    ext_model = model_for_task("drawing", model, user_override=user_override)
                    app.logger.info("PDF extract file=%s model=%s", file_name, ext_model)
                    if _is_gost_check_request(question, gost_mode):
                        gen = stream_extract_pdf_normative(
                            tmp_path,
                            file_name or "document.pdf",
                            question,
                            check_date=check_date,
                        )
                    else:
                        gen = stream_extract_pdf(
                            tmp_path, file_name or "document.pdf", question, ext_model, history
                        )
                    for chunk in gen:
                        if chunk.startswith("data: ") and chunk.strip() != "data: [DONE]":
                            try:
                                payload = json.loads(chunk[6:].strip())
                                if payload.get("error"):
                                    full_response.append(payload["error"])
                                else:
                                    full_response.append(payload.get("text", ""))
                            except Exception:
                                pass
                        yield chunk
                    extracted_text = "".join(full_response)

                elif ext in (".docx", ".doc"):
                    text = read_docx(tmp_path)
                    extracted_text = text
                    gen = (stream_text_response(text, question, model, history)
                           if estimate_tokens(text) <= profile["chunk_tokens"]
                           else stream_map_reduce(
                               [text[i:i+profile["chunk_tokens"]*4]
                                for i in range(0, len(text), profile["chunk_tokens"]*4)],
                               question, model, history))
                    for chunk in gen:
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".xlsx", ".xls", ".xlsm", ".ods"):
                    text = read_excel(tmp_path)
                    extracted_text = text
                    for chunk in stream_text_response(text, question, model, history):
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".csv", ".tsv"):
                    text = read_csv(tmp_path, ext)
                    extracted_text = text
                    for chunk in stream_text_response(text, question, model, history):
                        if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                            try:
                                full_response.append(json.loads(chunk[6:]).get("text", ""))
                            except Exception:
                                pass
                        yield chunk

                elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp", ".tiff", ".tif"):
                    if _is_gost_check_request(question, gost_mode):
                        for chunk in stream_extract_image_normative(
                            tmp_path,
                            file_name or "image.png",
                            question,
                            check_date=check_date,
                        ):
                            if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                try:
                                    full_response.append(json.loads(chunk[6:]).get("text", ""))
                                except Exception:
                                    pass
                            yield chunk
                    else:
                        with open(tmp_path, "rb") as f:
                            b64 = base64.b64encode(f.read()).decode("utf-8")
                            for chunk in stream_image_response([b64], question, model, history):
                                if chunk.startswith("data: ") and chunk != "data: [DONE]\n\n":
                                    try:
                                        full_response.append(json.loads(chunk[6:]).get("text", ""))
                                    except Exception:
                                        pass
                                yield chunk
                else:
                    yield f"data: {json.dumps({'error': f'Формат {ext} не поддерживается'})}\n\n"

        except Exception as e:
            app.logger.exception("chat generate")
            yield from _sse_error(f"Ошибка обработки: {_ollama_user_message(e)}")

        finally:
            if tmp_path:
                os.unlink(tmp_path)

            # Сохраняем ответ ассистента
            assistant_text = "".join(full_response)
            if assistant_text:
                save_message(conv_id, "assistant", assistant_text, model,
                             attachment_name=file_name,
                             attachment_text=extracted_text)

                # Генерируем название чата после первого ответа
                with get_db() as conn:
                    with conn.cursor() as cur:
                        cur.execute(
                            "SELECT COUNT(*) as cnt FROM messages WHERE conversation_id = %s AND role = 'user'",
                            (conv_id,)
                        )
                        cnt = cur.fetchone()["cnt"]

                if cnt == 1:
                    title = derive_chat_title(question, file_name)
                    with get_db() as conn:
                        with conn.cursor() as cur:
                            cur.execute("UPDATE conversations SET title = %s WHERE id = %s", (title, conv_id))
                        conn.commit()
                    yield f"data: {json.dumps({'title': title, 'conv_id': conv_id})}\n\n"

                threading.Thread(
                    target=maybe_summarize,
                    args=(conv_id, model, profile["num_ctx"]),
                    daemon=True
                ).start()

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


# ── статика ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory(ROOT_DIR, "index.html")


@app.route("/sso-login")
def sso_login():
    username = _normalize_username(request.args.get("u", ""))
    display_name = (request.args.get("d") or "").strip()
    ts_raw = (request.args.get("ts") or "").strip()
    sig = (request.args.get("sig") or "").strip()
    if not username or not display_name or not ts_raw or not sig:
        return "SSO parameters are missing", 400
    if not _verify_sso_payload(username, display_name, ts_raw, sig):
        return "SSO verification failed", 403
    session["logged_in"] = True
    session["username"] = username
    session["display_name"] = display_name
    return redirect("/")


@app.route("/api/me")
def api_me():
    username = _normalize_username(session.get("username") or "")
    display_name = (session.get("display_name") or username or "").strip()
    return jsonify({
        "logged_in": bool(session.get("logged_in")),
        "username": username,
        "display_name": display_name
    })


@app.route("/back.png")
def bg_image():
    return send_from_directory(ROOT_DIR, "back.png")

@app.route("/ico.png")
def ico_image():
    return send_from_directory(ROOT_DIR, "ico.png")

@app.route("/name.png")
def name_image():
    return send_from_directory(ROOT_DIR, "name.png")


@app.route("/api/models")
def api_models():
    available = get_available_models()
    result, seen_labels = [], set()
    for m in available:
        if not is_chat_model(m):
            continue
        label = MODEL_LABELS.get(m, m)
        if label not in seen_labels:
            result.append({"id": m, "label": label})
            seen_labels.add(label)
    if not result:
        result.append({"id": MODEL_DEFAULT, "label": MODEL_LABELS.get(MODEL_DEFAULT, MODEL_DEFAULT)})
    return {"models": result}


@app.route("/api/detect-file", methods=["POST"])
def api_detect_file():
    """Определяет тип файла и предлагает модель. Для PDF проверяет, скан это или текст."""
    file = request.files.get("file")
    if not file or not file.filename:
        return {"type": "unknown", "scanned": False, "model": MODEL_DEFAULT, "reason": ""}

    ext = os.path.splitext(file.filename)[1].lower()
    available = get_available_models()

    if ext != ".pdf":
        model, reason = suggest_model(ext, "", available)
        return {"type": ext.lstrip("."), "scanned": False, "model": model, "reason": reason}

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=ensure_upload_temp_dir())
    tmp_path = tmp.name
    tmp.close()
    try:
        file.save(tmp_path)
    except OSError as e:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        if getattr(e, "errno", None) == 28:
            return {"error": "Недостаточно места на диске для загрузки файла.", "type": "pdf", "scanned": False, "model": MODEL_DEFAULT, "reason": ""}, 507
        raise
    try:
        is_scanned = is_scanned_pdf(tmp_path) or detect_pdf_type(tmp_path) == "scanned"
        cfg = model_drawing() or model_scan()
        model = _resolve_config_model(cfg, available) if cfg else None
        if not model:
            model, _ = suggest_model(".txt", "", available)
        reason = f"PDF — полное извлечение текста (OCR); модель для уточнений: {model}"
        return {
            "type": "pdf",
            "scanned": is_scanned,
            "drawing": True,
            "model": model,
            "reason": reason,
        }
    finally:
        os.unlink(tmp_path)


@app.route("/api/suggest-model", methods=["POST"])
def api_suggest_model():
    data = request.get_json(force=True)
    file_ext = data.get("file_ext", "").lower()
    question = data.get("question", "")
    available = get_available_models()
    model, reason = suggest_model(file_ext, question, available)
    return {"model": model, "reason": reason}


if __name__ == "__main__":
    init_db()
    print("🚀 Запуск на http://localhost:5000")
    app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
